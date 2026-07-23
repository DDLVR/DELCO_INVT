"""Genera docs/DESGLOSE_DELCO_INVENTARIO.docx con diagramas EDT y flujos."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / 'docs'
IMG_DIR = OUT_DIR / '_diagramas_tmp'
OUT_DOCX = OUT_DIR / 'DESGLOSE_DELCO_INVENTARIO.docx'
OUT_DOCX_ALT = OUT_DIR / 'DESGLOSE_DELCO_INVENTARIO_v2.docx'

# Paleta Delco (azul institucional, sin morado genérico)
C_BG = (248, 250, 252)
C_ROOT = (26, 58, 92)
C_L1 = (44, 82, 130)
C_L2 = (56, 116, 168)
C_ACCENT = (34, 139, 99)
C_WARN = (180, 120, 40)
C_LINE = (120, 140, 160)
C_WHITE = (255, 255, 255)
C_TEXT = (30, 40, 50)


def _font(size: int, bold: bool = False):
    candidates = [
        'C:/Windows/Fonts/segoeuib.ttf' if bold else 'C:/Windows/Fonts/segoeui.ttf',
        'C:/Windows/Fonts/arialbd.ttf' if bold else 'C:/Windows/Fonts/arial.ttf',
        'C:/Windows/Fonts/calibrib.ttf' if bold else 'C:/Windows/Fonts/calibri.ttf',
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


def _rounded_box(draw, xy, fill, outline=None, radius=12, width=2):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def _center_text(draw, box, text, font, fill=C_WHITE):
    x0, y0, x1, y1 = box
    bbox = draw.textbbox((0, 0), text, font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    tx = x0 + (x1 - x0 - tw) / 2
    ty = y0 + (y1 - y0 - th) / 2 - 1
    draw.text((tx, ty), text, font=font, fill=fill)


def _multiline_center(draw, box, lines, font, fill=C_WHITE, gap=2, fonts=None):
    x0, y0, x1, y1 = box
    use_fonts = fonts if fonts else [font] * len(lines)
    heights = []
    widths = []
    for line, f in zip(lines, use_fonts):
        bbox = draw.textbbox((0, 0), line, font=f)
        widths.append(bbox[2] - bbox[0])
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + gap * (len(lines) - 1)
    ty = y0 + (y1 - y0 - total_h) / 2
    for line, tw, th, f in zip(lines, widths, heights, use_fonts):
        tx = x0 + (x1 - x0 - tw) / 2
        draw.text((tx, ty), line, font=f, fill=fill)
        ty += th + gap


def draw_edt_sistema(path: Path):
    """EDT: raíz → Operación / Gobierno → ámbitos → detalle con más texto."""
    w, h = 1720, 1020
    img = Image.new('RGB', (w, h), C_BG)
    draw = ImageDraw.Draw(img)
    f_title = _font(24, True)
    f_root = _font(18, True)
    f_l1 = _font(13, True)
    f_kid = _font(11, True)
    f_sub = _font(10)
    f_cap = _font(13, True)
    note = _font(12)

    draw.text((40, 22), 'Estructura de desglose del sistema', font=f_title, fill=C_ROOT)

    # Raíz
    root = (610, 70, 1110, 135)
    _rounded_box(draw, root, C_ROOT, radius=14)
    _center_text(draw, root, 'DELCO Inventario', f_root)
    rcx = (root[0] + root[2]) // 2

    # Bifurcación en dos ramas (barra por encima de los caps, sin atravesarlos)
    y_bar = 168
    draw.line([(rcx, root[3]), (rcx, y_bar)], fill=C_LINE, width=3)

    op_mid, gov_cap_mid = 530, 1420
    op_cap = (200, 188, 860, 235)
    gov_cap = (1200, 188, 1640, 235)
    spine_x = 1160  # a la izquierda de las cajas verdes (no las atraviesa)

    # Tramos horizontales solo hasta el centro de cada cap
    draw.line([(op_mid, y_bar), (gov_cap_mid, y_bar)], fill=C_LINE, width=3)
    draw.line([(op_mid, y_bar), (op_mid, op_cap[1])], fill=C_LINE, width=2)
    draw.line([(gov_cap_mid, y_bar), (gov_cap_mid, gov_cap[1])], fill=C_LINE, width=2)

    _rounded_box(draw, op_cap, C_L1, radius=10)
    _rounded_box(draw, gov_cap, C_ACCENT, radius=10)
    _center_text(draw, op_cap, 'Operación diaria', f_cap)
    _center_text(draw, gov_cap, 'Gobierno y control', f_cap)

    # Operación: 4 columnas con título + detalle en dos líneas
    ops = [
        (
            40, 290, 270, 360, 'Inventario',
            [
                ('Activos de campo', 'Medidores, SIM y módems'),
                ('Estado y ubicación', 'Disponible, instalado, bodega'),
                ('Kardex', 'Movimientos e historial'),
            ],
        ),
        (
            290, 290, 520, 360, 'Clientes',
            [
                ('Ficha del punto', 'Datos del cliente y serie'),
                ('Historial de proyectos', 'Cambios Actual / Reemplazado'),
                ('Restricciones', 'Bloqueo con justificación'),
            ],
        ),
        (
            540, 290, 770, 360, 'Órdenes de trabajo',
            [
                ('Ciclo de vida', 'Creación hasta cierre'),
                ('Asignación', 'Técnico responsable'),
                ('Validación', 'Revisión y cierre admin.'),
            ],
        ),
        (
            790, 290, 1020, 360, 'MoreApp',
            [
                ('Sincronización', 'Ingreso desde terreno'),
                ('Cola de revisión', 'Casos pendientes'),
                ('Alertas y bloqueos', 'Avisos operativos'),
            ],
        ),
    ]

    # Barra de distribución BAJO el cap de Operación (no sobre el borde del texto)
    y_fan = op_cap[3] + 18  # 253
    mids = [((o[0] + o[2]) // 2) for o in ops]
    draw.line([(op_mid, op_cap[3]), (op_mid, y_fan)], fill=C_LINE, width=2)
    draw.line([(mids[0], y_fan), (mids[-1], y_fan)], fill=C_LINE, width=2)

    for x0, y0, x1, y1, title, kids in ops:
        mx = (x0 + x1) // 2
        # conector hasta el tope del ámbito (luego la caja lo cubre si hubiera solape)
        draw.line([(mx, y_fan), (mx, y0)], fill=C_LINE, width=2)
        _rounded_box(draw, (x0, y0, x1, y1), C_L1, radius=12)
        _center_text(draw, (x0, y0, x1, y1), title, f_l1)

        # hijos: primero líneas solo en huecos, después cajas encima
        kid_boxes = []
        ky = y1 + 18
        for titulo, detalle in kids:
            box = (x0 + 6, ky, x1 - 6, ky + 58)
            kid_boxes.append((box, titulo, detalle))
            ky += 68

        # del ámbito al primer hijo (tope)
        draw.line([(mx, y1), (mx, kid_boxes[0][0][1])], fill=C_LINE, width=2)
        # entre hijos: solo el espacio entre fondo de uno y tope del siguiente
        for i in range(len(kid_boxes) - 1):
            bottom = kid_boxes[i][0][3]
            top_next = kid_boxes[i + 1][0][1]
            draw.line([(mx, bottom), (mx, top_next)], fill=C_LINE, width=2)

        for box, titulo, detalle in kid_boxes:
            _rounded_box(draw, box, C_L2, radius=8)
            _multiline_center(
                draw, box, [titulo, detalle], f_kid, gap=3, fonts=[f_kid, f_sub],
            )

    # Gobierno: espina a la IZQUIERDA de las cajas; stubs hasta el borde
    govs = [
        (1200, 290, 1640, 365, 'Acceso y roles', 'Perfiles, permisos y segregación de funciones'),
        (1200, 385, 1640, 460, 'Reportes', 'Exportes, filtros e indicadores de gestión'),
        (1200, 480, 1640, 555, 'Auditoría', 'Trazabilidad de cambios y consultas de control'),
        (1200, 575, 1640, 650, 'Soporte y catálogos', 'Maestros, parámetros y datos de apoyo'),
    ]
    # del cap hacia la espina (sin cruzar texto)
    stem_y = gov_cap[3] + 14
    first_cy = (govs[0][1] + govs[0][3]) // 2
    draw.line([(gov_cap_mid, gov_cap[3]), (gov_cap_mid, stem_y)], fill=C_LINE, width=2)
    draw.line([(spine_x, stem_y), (gov_cap_mid, stem_y)], fill=C_LINE, width=2)
    draw.line([(spine_x, stem_y), (spine_x, first_cy)], fill=C_LINE, width=2)

    for i, (x0, y0, x1, y1, title, desc) in enumerate(govs):
        cy = (y0 + y1) // 2
        if i > 0:
            prev_bottom = govs[i - 1][3]
            # tramo vertical solo en el hueco entre cajas
            draw.line([(spine_x, prev_bottom), (spine_x, y0)], fill=C_LINE, width=2)
            draw.line([(spine_x, y0), (spine_x, cy)], fill=C_LINE, width=2)
        draw.line([(spine_x, cy), (x0, cy)], fill=C_LINE, width=2)
        _rounded_box(draw, (x0, y0, x1, y1), C_ACCENT, radius=12)
        _multiline_center(
            draw, (x0, y0, x1, y1), [title, desc], f_l1, gap=4, fonts=[f_l1, f_sub],
        )

    draw.text(
        (40, 960),
        'La EDT organiza el alcance en dos ramas: operación diaria (izquierda) y gobierno/control (derecha), '
        'sobre una misma plataforma de información.',
        font=note,
        fill=C_TEXT,
    )
    img.save(path)


def draw_flujo_operativo(path: Path):
    """Flujo punta a punta: terreno → sistema → cierre."""
    w, h = 1400, 520
    img = Image.new('RGB', (w, h), C_BG)
    draw = ImageDraw.Draw(img)
    f_title = _font(22, True)
    f_box = _font(14, True)
    f_small = _font(12)

    draw.text((40, 24), 'Flujo operativo — de terreno a decisión', font=f_title, fill=C_ROOT)

    stages = [
        (60, 120, 280, 220, C_ACCENT, ['Terreno', 'MoreApp']),
        (360, 120, 620, 220, C_L1, ['Sincronización', 'ingreso al sistema']),
        (700, 80, 980, 160, C_L2, ['Inventario', 'movimientos']),
        (700, 180, 980, 260, C_L2, ['Clientes', 'cruces / alertas']),
        (700, 280, 980, 360, C_L2, ['OT', 'vinculación']),
        (1060, 120, 1340, 220, C_WARN, ['Validación', 'administrativa']),
        (1060, 280, 1340, 380, C_ROOT, ['Reportes', 'y Auditoría']),
    ]
    for box in stages:
        x0, y0, x1, y1, color, lines = box
        _rounded_box(draw, (x0, y0, x1, y1), color, radius=14)
        _multiline_center(draw, (x0, y0, x1, y1), lines, f_box)

    # flechas
    def arrow(a, b):
        draw.line([a, b], fill=C_LINE, width=4)
        # punta simple
        bx, by = b
        draw.polygon([(bx, by), (bx - 10, by - 7), (bx - 10, by + 7)], fill=C_LINE)

    arrow((280, 170), (360, 170))
    arrow((620, 170), (700, 120))
    draw.line([(660, 170), (660, 320)], fill=C_LINE, width=3)
    draw.line([(660, 120), (700, 120)], fill=C_LINE, width=3)
    draw.line([(660, 220), (700, 220)], fill=C_LINE, width=3)
    draw.line([(660, 320), (700, 320)], fill=C_LINE, width=3)
    arrow((980, 200), (1060, 170))
    draw.line([(1020, 200), (1020, 330)], fill=C_LINE, width=3)
    arrow((1020, 330), (1060, 330))

    draw.text(
        (40, 450),
        'Una misma visita de terreno actualiza equipos, puede alertar sobre el cliente y alimentar la OT; luego se valida y se reporta.',
        font=f_small,
        fill=C_TEXT,
    )
    img.save(path)


def draw_flujo_ot(path: Path):
    """Ciclo de vida OT horizontal."""
    w, h = 1400, 360
    img = Image.new('RGB', (w, h), C_BG)
    draw = ImageDraw.Draw(img)
    f_title = _font(22, True)
    f_box = _font(13, True)
    f_small = _font(12)

    draw.text((40, 24), 'Flujo de orden de trabajo (OT)', font=f_title, fill=C_ROOT)

    steps = [
        'Creada',
        'Asignada',
        'En ejecución',
        'Informada\n(terreno)',
        'Validada',
        'Finalizada',
    ]
    colors = [C_L2, C_L1, C_L1, C_ACCENT, C_WARN, C_ROOT]
    x = 40
    y0, y1 = 120, 210
    boxes = []
    for i, (label, color) in enumerate(zip(steps, colors)):
        box = (x, y0, x + 180, y1)
        boxes.append(box)
        _rounded_box(draw, box, color, radius=12)
        _multiline_center(draw, box, label.split('\n'), f_box)
        if i < len(steps) - 1:
            ax0 = x + 180
            ax1 = x + 210
            mid = (y0 + y1) // 2
            draw.line([(ax0, mid), (ax1, mid)], fill=C_LINE, width=4)
            draw.polygon([(ax1, mid), (ax1 - 10, mid - 7), (ax1 - 10, mid + 7)], fill=C_LINE)
        x += 210

    # rama cancelada
    _rounded_box(draw, (850, 260, 1090, 320), (140, 70, 70), radius=10)
    _center_text(draw, (850, 260, 1090, 320), 'Cancelada (salida)', f_box)
    draw.line([(930, 210), (930, 260)], fill=C_LINE, width=3)

    draw.text(
        (40, 300),
        'La OT ordena el trabajo; MoreApp aporta evidencia de terreno; Administración cierra el ciclo con validación.',
        font=f_small,
        fill=C_TEXT,
    )
    img.save(path)


def draw_flujo_moreapp(path: Path):
    """Flujo MoreApp: sync → cruce → revisión."""
    w, h = 1400, 420
    img = Image.new('RGB', (w, h), C_BG)
    draw = ImageDraw.Draw(img)
    f_title = _font(22, True)
    f_box = _font(13, True)
    f_small = _font(12)

    draw.text((40, 24), 'Flujo MoreApp — ingreso y revisión', font=f_title, fill=C_ROOT)

    boxes = [
        (60, 110, 300, 200, C_ACCENT, ['Formulario', 'en terreno']),
        (380, 110, 640, 200, C_L1, ['Sync / webhook', 'al servidor']),
        (720, 60, 980, 140, C_L2, ['Cruce con', 'inventario']),
        (720, 160, 980, 240, C_L2, ['Cruce con', 'cliente / OT']),
        (1060, 60, 1340, 140, C_WARN, ['Advertencia', 'o bloqueo']),
        (1060, 160, 1340, 240, C_ROOT, ['Revisado /', 'Descartado']),
    ]
    for x0, y0, x1, y1, color, lines in boxes:
        _rounded_box(draw, (x0, y0, x1, y1), color, radius=12)
        _multiline_center(draw, (x0, y0, x1, y1), lines, f_box)

    draw.line([(300, 155), (380, 155)], fill=C_LINE, width=4)
    draw.polygon([(380, 155), (370, 148), (370, 162)], fill=C_LINE)
    draw.line([(640, 155), (680, 155), (680, 100), (720, 100)], fill=C_LINE, width=3)
    draw.line([(680, 155), (680, 200), (720, 200)], fill=C_LINE, width=3)
    draw.line([(980, 100), (1060, 100)], fill=C_LINE, width=4)
    draw.line([(980, 200), (1060, 200)], fill=C_LINE, width=4)

    _rounded_box(draw, (380, 280, 980, 360), C_WHITE, outline=C_L1, radius=12, width=3)
    _multiline_center(
        draw,
        (380, 280, 980, 360),
        ['Cola operativa: Pendiente  →  Con advertencia  →  Revisado / Descartado'],
        f_box,
        fill=C_ROOT,
    )

    draw.text(
        (40, 380),
        'Lo automático no reemplaza el criterio: los bloqueos y alertas quedan visibles para revisión humana.',
        font=f_small,
        fill=C_TEXT,
    )
    img.save(path)


def draw_edt_roles(path: Path):
    """EDT ligero de roles / quién usa qué."""
    w, h = 1400, 480
    img = Image.new('RGB', (w, h), C_BG)
    draw = ImageDraw.Draw(img)
    f_title = _font(22, True)
    f_box = _font(14, True)
    f_small = _font(12)

    draw.text((40, 24), 'Responsabilidades por rol', font=f_title, fill=C_ROOT)

    roles = [
        (50, 100, 300, 200, C_ROOT, ['ADMIN', 'Gobierno total']),
        (330, 100, 580, 200, C_L1, ['ADMINISTRATIVO', 'Operación diaria']),
        (610, 100, 860, 200, C_ACCENT, ['TÉCNICO', 'Trabajo en terreno']),
        (890, 100, 1140, 200, C_WARN, ['GERENCIA', 'Visión agregada']),
        (1170, 100, 1360, 200, C_L2, ['AUDITOR', 'Trazabilidad']),
    ]
    for x0, y0, x1, y1, color, lines in roles:
        _rounded_box(draw, (x0, y0, x1, y1), color, radius=12)
        _multiline_center(draw, (x0, y0, x1, y1), lines, f_box)

    # foco debajo
    focos = [
        (50, 260, 300, 400, ['Usuarios', 'configuración', 'control']),
        (330, 260, 580, 400, ['Clientes', 'OT', 'MoreApp', 'inventario']),
        (610, 260, 860, 400, ['Sus OT', 'evidencia', 'terreno']),
        (890, 260, 1140, 400, ['Dashboards', 'indicadores']),
        (1170, 260, 1360, 400, ['Auditoría', 'consulta']),
    ]
    for x0, y0, x1, y1, lines in focos:
        _rounded_box(draw, (x0, y0, x1, y1), C_WHITE, outline=C_LINE, radius=12, width=2)
        _multiline_center(draw, (x0, y0, x1, y1), lines, f_small, fill=C_TEXT, gap=4)
        mx = (x0 + x1) // 2
        draw.line([(mx, 200), (mx, 260)], fill=C_LINE, width=3)

    img.save(path)


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True, color=(26, 58, 92))
    return p


def add_para(doc, text, size=11, bold=False, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=(40, 50, 60))
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(2)
    return p


def add_image(doc, path: Path, width_cm=16.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    p.paragraph_format.space_after = Pt(12)


def build_docx(images: dict):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Portada
    add_para(doc, 'DELCO Chile', size=12, bold=True, center=True)
    add_para(doc, 'DELCO Inventario', size=26, bold=True, center=True)
    add_para(doc, 'Desglose operativo del sistema', size=14, center=True)
    add_para(doc, 'Julio 2026', size=10, center=True)
    add_para(
        doc,
        'Plataforma unificada de inventario, clientes, órdenes de trabajo e integración '
        'con formularios de terreno (MoreApp), con trazabilidad y control de gestión.',
        size=11,
        center=True,
    )
    add_para(
        doc,
        'Este documento describe la estructura funcional del sistema y los flujos que '
        'conectan el trabajo en terreno con la operación de oficina, la validación '
        'administrativa y la mirada de control.',
        size=11,
        center=True,
    )

    doc.add_page_break()

    add_heading(doc, 'Estructura de desglose', 1)
    add_para(
        doc,
        'El sistema se entiende como una estructura de desglose (EDT): un mapa del alcance '
        'funcional organizado en dos ramas que comparten la misma base de información. '
        'La rama de operación diaria concentra lo que ocurre en el día a día —inventario, '
        'clientes, órdenes de trabajo e ingreso desde MoreApp—. La rama de gobierno y '
        'control concentra lo que ordena, protege y explota esa información: acceso y '
        'roles, reportes, auditoría y catálogos de soporte.',
    )
    add_para(
        doc,
        'Esta organización permite presentar el alcance sin depender de un recorrido '
        'pantalla por pantalla. Cada bloque del diagrama corresponde a una capacidad '
        'con responsable, datos propios y reglas de negocio. El valor de la EDT es '
        'mostrar que inventario, cliente y orden de trabajo no son módulos aislados, '
        'sino partes de un mismo ciclo operativo que luego se gobierna con reportes '
        'y trazabilidad.',
    )
    add_image(doc, images['edt'], width_cm=17)
    add_para(
        doc,
        'En operación, Inventario administra activos de campo (medidores, SIM y módems), '
        'su estado y ubicación, y el kardex de movimientos. Clientes concentra la ficha '
        'del punto, el historial de proyectos (Actual / Reemplazado) y las restricciones '
        'con justificación. Órdenes de trabajo cubre el ciclo de vida, la asignación al '
        'técnico y la validación de cierre. MoreApp aporta la sincronización desde '
        'terreno, la cola de revisión y las alertas o bloqueos operativos.',
    )
    add_para(
        doc,
        'En gobierno y control, Acceso y roles define quién ingresa y con qué permisos; '
        'Reportes permite filtrar y exportar información de gestión; Auditoría conserva '
        'trazabilidad de cambios relevantes; y Soporte y catálogos mantienen maestros '
        'y parámetros que sostienen al resto del sistema.',
    )

    add_heading(doc, 'Responsabilidades por rol', 1)
    add_para(
        doc,
        'El acceso está diferenciado por perfil. Administración y operación diaria '
        'concentran la gestión de fichas, equipos e informes; el técnico trabaja sobre '
        'sus órdenes; gerencia observa indicadores; auditoría consulta la trazabilidad '
        'de cambios relevantes.',
    )
    add_image(doc, images['roles'])

    doc.add_page_break()

    add_heading(doc, 'Flujo operativo', 1)
    add_para(
        doc,
        'El flujo punta a punta muestra cómo una actividad de terreno se transforma en '
        'información operativa. Al sincronizar MoreApp, el sistema puede actualizar '
        'inventario, cruzar datos de cliente y vincular evidencia a la orden de trabajo. '
        'Luego interviene la validación administrativa y, finalmente, la explotación '
        'mediante reportes y auditoría.',
    )
    add_image(doc, images['flujo'])

    add_heading(doc, 'Ciclo de la orden de trabajo', 1)
    add_para(
        doc,
        'La orden de trabajo es el eje de ejecución. Avanza desde su creación y '
        'asignación hasta la ejecución en terreno, la incorporación de evidencia y el '
        'cierre por validación. Cuando corresponde, también contempla la cancelación '
        'como salida controlada del ciclo.',
    )
    add_image(doc, images['ot'])

    add_heading(doc, 'Integración MoreApp', 1)
    add_para(
        doc,
        'MoreApp aporta el canal de captura en terreno. La sincronización reduce carga '
        'manual, pero no elimina la revisión: advertencias y bloqueos operativos quedan '
        'visibles en cola para que el equipo administrativo resuelva con criterio, '
        'manteniendo consistencia entre lo reportado en campo y el maestro del sistema.',
    )
    add_image(doc, images['moreapp'])

    try:
        doc.save(OUT_DOCX)
        print(f'Generado: {OUT_DOCX}')
    except PermissionError:
        doc.save(OUT_DOCX_ALT)
        print(f'Archivo en uso; generado: {OUT_DOCX_ALT}')


def main():
    IMG_DIR.mkdir(parents=True, exist_ok=True)
    images = {
        'edt': IMG_DIR / 'edt_sistema.png',
        'roles': IMG_DIR / 'edt_roles.png',
        'flujo': IMG_DIR / 'flujo_operativo.png',
        'ot': IMG_DIR / 'flujo_ot.png',
        'moreapp': IMG_DIR / 'flujo_moreapp.png',
    }
    draw_edt_sistema(images['edt'])
    draw_edt_roles(images['roles'])
    draw_flujo_operativo(images['flujo'])
    draw_flujo_ot(images['ot'])
    draw_flujo_moreapp(images['moreapp'])
    build_docx(images)


if __name__ == '__main__':
    main()
